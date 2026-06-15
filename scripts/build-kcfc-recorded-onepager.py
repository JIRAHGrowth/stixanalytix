"""Build the recorded-and-emailed version of the KCFC GK Parent 1-pager."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"C:\Users\joshu\OneDrive\KCITY FC\04 - Documents\Communications\KCFC_Seven_Messages_Recorded_Video_and_Email.docx")

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_para(text="", *, bold=False, italic=False, size=None, space_after=6, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if not text:
        return p
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_runs(runs, *, space_after=6, align=None):
    """runs = list of (text, dict_of_attrs). attrs: bold, italic, size, color."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for text, attrs in runs:
        r = p.add_run(text)
        r.bold = attrs.get("bold", False)
        r.italic = attrs.get("italic", False)
        if attrs.get("size"):
            r.font.size = Pt(attrs["size"])
        if attrs.get("color"):
            r.font.color.rgb = attrs["color"]
    return p


def add_speaker(name, line):
    add_runs([(f"{name}:  ", {"bold": True}), (line, {})], space_after=8)


def add_cue(text):
    add_para(f"[{text}]", italic=True, size=10, space_after=8, color=RGBColor(0x6B, 0x6B, 0x6B))


def add_h1(text):
    add_para(text, bold=True, size=20, space_after=4)


def add_h2(text):
    add_para(text, bold=True, size=14, space_after=4)


def add_h3(text):
    add_para(text, bold=True, size=11, space_after=4)


def add_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------- Header ----------
add_para("KELOWNA CITY FC", bold=True, size=10, space_after=2,
         color=RGBColor(0x6B, 0x6B, 0x6B))
add_h1("Seven Messages — Recorded Video + Email")
add_para("The same seven messages, adapted from the in-person script for a recorded delivery and email distribution.",
         italic=True, size=10, space_after=14, color=RGBColor(0x6B, 0x6B, 0x6B))

# ---------- SECTION A: VIDEO SCRIPT ----------
add_h2("Section A — Video script")
add_para("Target length: 5–6 minutes. Speakers: Joshua + Ezra. Read in your own cadence; the words are a floor, not a ceiling.",
         italic=True, size=10, space_after=12, color=RGBColor(0x6B, 0x6B, 0x6B))

add_h3("Open")
add_speaker("Joshua",
            "Hi — I'm Joshua, the goalkeeper coach at Kelowna City FC. We wanted every "
            "goalkeeper family in one room for this conversation, but the schedules didn't "
            "line up for everyone. Rather than cut people out, Ezra and I recorded it. Same "
            "conversation we would have had with you in person.")
add_speaker("Ezra",
            "I'm Ezra. I run the club. We'll take about five minutes. There's also an email "
            "underneath this video with everything written down, so you don't have to "
            "remember any of it.")
add_cue("beat — transition")

add_h3("1.  We are building a pyramid your goalkeeper has a place in for life.")
add_speaker("Joshua",
            "First — we're building a pyramid your goalkeeper has a place in for life. "
            "Three to four hundred minis at the base. Around 140 goalkeepers across 60 to 80 "
            "teams in the middle. A BC Premier League team at the top, and adult divisions "
            "beyond that. This isn't aspiration. It's a deliberate pathway. Your goalkeeper "
            "enters once and has somewhere to keep going.")
add_speaker("Ezra",
            "That's the frame. Everything we're about to say sits underneath it.")

add_h3("2.  Goalkeepers are structural to that pyramid, not a side project.")
add_speaker("Joshua",
            "Second — goalkeepers aren't a side project inside that pyramid. They're "
            "structural to it. The club needs 140 of them for the model to work — two on "
            "every team. That's why, starting this season, goalkeeping has its own coach, "
            "its own structure, and yes, its own fee.")

add_h3("3.  Two changes are landing together, and we are owning it openly.")
add_speaker("Ezra",
            "Third — and we want to be straight with you on this — two things are changing "
            "at once. The base fee went up. The goalkeeper fee is new. We're not going to "
            "pretend either of those is small. We owe you the reasoning on both, and that's "
            "the next two points.")

add_h3("4.  We are showing you the math.")
add_speaker("Ezra",
            "So here is the math, in plain terms. What the program actually costs to run. "
            "What family fees cover. Where sponsorship comes in. And the gap the club is "
            "carrying on purpose — because we'd rather build this program well than charge "
            "you harder for it. The numbers are in the email below, so you can sit with them.")
add_speaker("Joshua",
            "I'll walk those numbers in the email. Take your time with them.")

add_h3("5.  Your goalkeeper is getting something genuinely different.")
add_speaker("Joshua",
            "Fifth — and this is the one I want you to take seriously — your goalkeeper is "
            "getting something genuinely different this year. Roughly 100 opportunities to "
            "train the position. Team training included on top of that. Personalized kit. "
            "Written progress reports you can actually use. A new coach joining the program. "
            "And per session, it's less than any local clinic you'd find. We can defend that "
            "number.")

add_h3("6.  The founding rate is for this season only, and we are telling you now.")
add_speaker("Ezra",
            "Sixth — the rate this season is a founding rate, and we want you to know that "
            "now, not later. $399 this year. $599 next year. $799 standard from year three "
            "on. We're saying it up front so it never feels like a surprise.")

add_h3("7.  No goalkeeper is kept out because of cost.")
add_speaker("Ezra",
            "And the last one is the most important one. No goalkeeper is kept out of this "
            "program because of cost. If money is a real obstacle for your family, come find "
            "me. It is a normal conversation. Reply to the email this video came in, or text "
            "me directly at [phone]. We figure it out. That's the rule.")
add_cue("long pause — let the seventh sit before closing")

add_h3("Close")
add_speaker("Joshua",
            "That's the seven. No decision needed today. Registration opens [date]. The "
            "founding-rate window closes [date].")
add_speaker("Ezra",
            "If something didn't land, or there's a question we should have answered, just "
            "reply to the email. We'll get back to you.")
add_speaker("Joshua",
            "Thanks for the five minutes.")

add_rule()

# ---------- SECTION B: EMAIL BODY ----------
add_h2("Section B — Email body")
add_para("Send this with the video embedded or linked at the top. The seven messages are repeated below the video so parents have something durable to refer back to.",
         italic=True, size=10, space_after=12, color=RGBColor(0x6B, 0x6B, 0x6B))

add_runs([("Subject:  ", {"bold": True}),
          ("Seven things about the goalkeeper program — a short video from Ezra and Joshua", {})],
         space_after=12)

add_para("Hi [first name],", space_after=8)
add_para("Ezra and I wanted to get every goalkeeper family in one room to talk through what's "
         "changing for the goalkeeper program this season. Schedules didn't line up for "
         "everyone, so we recorded it instead. Five minutes. Same conversation we would have "
         "had with you in person.", space_after=10)

add_para("▶  [VIDEO LINK]", bold=True, space_after=14)

add_para("Below are the same seven messages, in writing, so you have them to refer back to.",
         space_after=12)

add_para("1.  We are building a pyramid your goalkeeper has a place in for life.", bold=True, space_after=4)
add_para("300–400 minis at the base. ~140 goalkeepers across 60–80 teams. A BC Premier League "
         "team at the top. Adult divisions beyond that. Your goalkeeper enters the pyramid "
         "once and has somewhere to keep going.", space_after=12)

add_para("2.  Goalkeepers are structural to that pyramid, not a side project.", bold=True, space_after=4)
add_para("The club needs ~140 goalkeepers to function as designed — two on every team. That's "
         "why goalkeeping now has its own coach, its own structure, and its own fee.",
         space_after=12)

add_para("3.  Two changes are landing together, and we are owning it openly.", bold=True, space_after=4)
add_para("The base fee went up. The goalkeeper fee is new. Neither is small, and we would "
         "rather explain than soften.", space_after=12)

add_para("4.  Here is the math.", bold=True, space_after=4)
add_para("•  Program cost:  [$X]", space_after=2)
add_para("•  Covered by family fees:  [$X]", space_after=2)
add_para("•  Covered by sponsorship:  [$X]", space_after=2)
add_para("•  Carried by the club, on purpose:  [$X]", space_after=6)
add_para("We would rather build this program well than charge harder for it.", space_after=12)

add_para("5.  Your goalkeeper is getting something genuinely different.", bold=True, space_after=4)
add_para("•  ~100 opportunities to train the position this year", space_after=2)
add_para("•  Team training included on top", space_after=2)
add_para("•  Personalized kit", space_after=2)
add_para("•  Written progress reports", space_after=2)
add_para("•  A new coach joining the program", space_after=2)
add_para("•  Less per session than any local clinic", space_after=12)

add_para("6.  The founding rate is for this season only.", bold=True, space_after=4)
add_para("•  $399 this year", space_after=2)
add_para("•  $599 next year", space_after=2)
add_para("•  $799 standard from year three on", space_after=6)
add_para("We are telling you up front so it never becomes a surprise later.", space_after=12)

add_para("7.  No goalkeeper is kept out because of cost.", bold=True, space_after=4)
add_para("If money is a real obstacle, reply to this email or text Ezra directly at [phone]. "
         "It is a normal conversation. That part is non-negotiable for us.", space_after=14)

add_para("Key dates", bold=True, space_after=4)
add_para("•  Registration opens:  [date]", space_after=2)
add_para("•  Founding-rate window closes:  [date]", space_after=12)

add_para("No decision needed today. If something didn't land or you have a question we should "
         "have answered, just reply.", space_after=14)

add_para("Joshua Marshall  —  Goalkeeper Coach", space_after=2)
add_para("Ezra [last name]  —  [title]", space_after=2)
add_para("Kelowna City FC", space_after=12)

add_rule()

# ---------- SECTION C: PRODUCTION NOTES ----------
add_h2("Section C — Production notes")
add_para("Small choices that prevent a fourth re-shoot.",
         italic=True, size=10, space_after=10, color=RGBColor(0x6B, 0x6B, 0x6B))

add_para("Length", bold=True, space_after=4)
add_para("5–6 minutes is the target. Past 7 and parents stop watching. If you run long, cut "
         "from messages 1, 2, 4 — keep 3, 5, 6, 7 verbatim. Message 7 is the one that has to "
         "land cleanly.", space_after=10)

add_para("Framing", bold=True, space_after=4)
add_para("Two-shot with both of you in frame is ideal — it visually says \"we are aligned.\" "
         "If that isn't workable, alternate medium close-ups; speaker on screen leads, the "
         "other is named in the email caption. Phone in landscape, tripod or stable surface, "
         "eye-line just above the lens.", space_after=10)

add_para("Light + sound", bold=True, space_after=4)
add_para("Soft front light (a window during the day is enough). Quiet room — bad audio loses "
         "parents faster than bad video. Record one test minute and play it back on phone "
         "speakers before committing to a take.", space_after=10)

add_para("Takes", bold=True, space_after=4)
add_para("Read the whole script through together once, off-camera. Then record in 1–3 takes. "
         "Don't chase perfection — warmth and clarity beat polish.", space_after=10)

add_para("Captions", bold=True, space_after=4)
add_para("Add captions before sending. Many parents will watch with the sound off (on a "
         "phone, between work and pickup). Captions roughly double watch-through.",
         space_after=10)

add_para("Distribution", bold=True, space_after=4)
add_para("Upload to YouTube as Unlisted (not Private — Unlisted lets the link work without "
         "sign-in) and embed or link in the email. Avoid attaching the video file directly; "
         "it will bounce on some inboxes.", space_after=10)

add_para("Reply path", bold=True, space_after=4)
add_para("Make sure the email is sent from Ezra's address (or one he checks daily), because "
         "message 7 routes replies to him. Joshua copied so you both see questions in real "
         "time.", space_after=8)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Wrote: {OUT}")
